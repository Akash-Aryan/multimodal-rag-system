"""Document processing module for text documents."""

from typing import Dict, List, Any, Union
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib
from datetime import datetime

import PyPDF2
import pdfplumber
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from loguru import logger

from ..config import RAGConfig


class DocumentProcessor:
    """Process text documents (PDF, DOCX, TXT, MD)."""
    
    def __init__(self, config: RAGConfig):
        """Initialize document processor."""
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.document_processing.chunk_size,
            chunk_overlap=config.document_processing.chunk_overlap,
            separators=["\\n\\n", "\\n", ". ", " ", ""]
        )
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def process_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Process a document file and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing processed content and metadata
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text_content = await self._process_pdf(file_path)
        elif file_extension == '.docx':
            text_content = await self._process_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            text_content = await self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported document format: {file_extension}")
        
        # Split text into chunks
        chunks = await self._split_text(text_content)
        
        # Generate metadata
        file_hash = await self._calculate_file_hash(file_path)
        
        result = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_type': 'document',
            'file_extension': file_extension.lstrip('.'),
            'file_hash': file_hash,
            'file_size': file_path.stat().st_size,
            'processed_at': datetime.now().isoformat(),
            'total_length': len(text_content),
            'num_chunks': len(chunks),
            'content': text_content,
            'chunks': [
                {
                    'chunk_id': i,
                    'text': chunk,
                    'length': len(chunk),
                    'metadata': {
                        'source_file': str(file_path),
                        'chunk_index': i,
                        'file_type': 'document'
                    }
                }
                for i, chunk in enumerate(chunks)
            ]
        }
        
        logger.info(f"Processed document {file_path}: {len(chunks)} chunks")
        return result
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        def extract_text():
            text = ""
            
            # Try with pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\\n"
                
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\\n"
                return text.strip()
            except Exception as e:
                logger.error(f"Failed to extract text from PDF {file_path}: {e}")
                return ""
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, extract_text
        )
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        def extract_text():
            try:
                doc = Document(file_path)
                text = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\\n"
                
                return text.strip()
            except Exception as e:
                logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
                return ""
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, extract_text
        )
    
    async def _process_text(self, file_path: Path) -> str:
        """Read plain text file."""
        def read_text():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                # Try with different encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            return file.read()
                    except UnicodeDecodeError:
                        continue
                
                logger.error(f"Failed to read text file {file_path} with any encoding")
                return ""
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, read_text
        )
    
    async def _split_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        def split():
            return self.text_splitter.split_text(text)
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, split
        )
    
    async def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file."""
        def calculate_hash():
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, calculate_hash
        )
